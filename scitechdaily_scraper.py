"""Export SciTechDaily articles in a date range to one DOCX per article."""

from __future__ import annotations

import argparse
import os
import re
import time
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Callable
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


BASE_URL = "https://scitechdaily.com/"
INVALID_FILENAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--start", required=True, type=date.fromisoformat)
    parser.add_argument("--end", default=date.today().isoformat(), type=date.fromisoformat)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--delay", type=float, default=0.2)
    return parser.parse_args()


def make_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(total=4, backoff_factor=1, status_forcelist=(429, 500, 502, 503, 504))
    session.mount("https://", HTTPAdapter(max_retries=retry))
    session.headers["User-Agent"] = "Mozilla/5.0 (compatible; research document exporter/1.0)"
    return session


def archive_articles(session: requests.Session, start: date, end: date) -> list[tuple[date, str, str]]:
    found: dict[str, tuple[date, str, str]] = {}
    for page_no in range(1, 100):
        url = BASE_URL if page_no == 1 else urljoin(BASE_URL, f"page/{page_no}/")
        response = session.get(url, timeout=45)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        page_dates: list[date] = []
        for article in soup.select("article"):
            time_tag = article.select_one("time[datetime]")
            title_link = article.select_one("h2 a[href], h3 a[href]")
            if not time_tag or not title_link:
                continue
            published = datetime.fromisoformat(time_tag["datetime"]).date()
            page_dates.append(published)
            if start <= published <= end:
                article_url = urljoin(BASE_URL, title_link["href"])
                found[article_url] = (published, title_link.get_text(" ", strip=True), article_url)
        print(f"Archive page {page_no}: {len(page_dates)} dated entries")
        if page_dates and min(page_dates) < start:
            break
    return sorted(found.values(), key=lambda row: (row[0], row[1]), reverse=True)


def safe_filename(title: str, used: set[str], folder: Path | None = None) -> str:
    """Unique .docx name, budgeted so the full path stays under Windows MAX_PATH."""
    stem = INVALID_FILENAME.sub("", title).rstrip(" .") or "Untitled article"
    budget = 170
    if folder is not None:
        budget = 250 - len(str(folder)) - len(".docx") - len(os.sep) - len(" (99)")
        budget = max(24, min(170, budget))
    stem = stem[:budget].rstrip(" .") or "Untitled article"
    candidate = stem
    number = 2
    while candidate.casefold() in used:
        suffix = f" ({number})"
        candidate = stem[:budget - len(suffix)].rstrip(" .") + suffix
        number += 1
    used.add(candidate.casefold())
    return candidate + ".docx"


def image_url(img: Tag, page_url: str) -> str | None:
    value = img.get("data-lazy-src") or img.get("data-src") or img.get("src")
    if not value or value.startswith("data:"):
        return None
    return urljoin(page_url, value)


def add_image(document: Document, session: requests.Session, img: Tag, page_url: str) -> bool:
    url = image_url(img, page_url)
    if not url:
        return False
    try:
        response = session.get(url, timeout=45)
        response.raise_for_status()
        raw = BytesIO(response.content)
        with Image.open(raw) as parsed:
            parsed.verify()
        raw.seek(0)
        document.add_picture(raw, width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception as exc:
        document.add_paragraph(f"[Image unavailable: {url}]")
        print(f"  Image warning: {exc}")
        return False


def write_article(session: requests.Session, entry: tuple[date, str, str], target: Path,
                  screen: Callable[[str, str], None] | None = None,
                  classify: Callable[[str, str], bool] | None = None,
                  news_dir: Path | None = None) -> tuple[int, str, str]:
    """Render one article to DOCX. Returns (image count, resolved title, category).

    `screen` is called with (title, body text) before any image is downloaded; it
    may raise to reject the article. `classify(title, body)` (optional) decides the
    category: a False result ('news') routes the file to `news_dir`. news_scraper
    passes its keyword/duplicate check and its relevance classifier through here so
    behaviour is identical across all sites.
    """
    _published, listed_title, url = entry
    response = session.get(url, timeout=45)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    title_tag = soup.select_one("article h1, h1.entry-title")
    content = soup.select_one("article .entry-content, .post-content.entry-content")
    if not content:
        raise RuntimeError("Article body was not found")

    title = title_tag.get_text(" ", strip=True) if title_tag else listed_title
    body_text = " ".join(content.get_text(" ", strip=True).split())
    if screen is not None:
        screen(title, body_text)
    category = "technical"
    if classify is not None and not classify(title, body_text):
        category = "news"
        if news_dir is not None:
            target = news_dir / target.name

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    document.add_heading(title, level=0)

    image_count = 0
    emitted_images: set[str] = set()
    for node in content.find_all(["figure", "p", "h2", "h3", "h4", "h5", "blockquote", "ul", "ol"], recursive=True):
        if node.find_parent(["figure", "p", "blockquote", "ul", "ol"]) is not None:
            continue
        classes = set(node.get("class", []))
        if classes & {"newsletter-inline", "sharedaddy", "post-tags", "code-block"}:
            continue
        if node.name == "figure":
            for img in node.find_all("img"):
                url_value = image_url(img, url)
                if url_value and url_value not in emitted_images:
                    image_count += int(add_image(document, session, img, url))
                    emitted_images.add(url_value)
            caption = node.find("figcaption") or node.select_one(".wp-caption-text")
            if caption and caption.get_text(" ", strip=True):
                paragraph = document.add_paragraph(caption.get_text(" ", strip=True))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].italic = True
            continue
        for img in node.find_all("img"):
            url_value = image_url(img, url)
            if url_value and url_value not in emitted_images:
                image_count += int(add_image(document, session, img, url))
                emitted_images.add(url_value)
        text = " ".join(node.get_text(" ", strip=True).split())
        if not text or text.startswith("Never miss a breakthrough:") or text.startswith("Follow us on Google"):
            continue
        if node.name in {"h2", "h3", "h4", "h5"}:
            document.add_heading(text, level=1 if node.name in {"h2", "h3"} else 2)
        elif node.name == "blockquote":
            document.add_paragraph(text, style="Quote")
        elif node.name in {"ul", "ol"}:
            style = "List Bullet" if node.name == "ul" else "List Number"
            for item in node.find_all("li", recursive=False):
                document.add_paragraph(" ".join(item.get_text(" ", strip=True).split()), style=style)
        else:
            document.add_paragraph(text)

    props = document.core_properties
    props.title = title
    props.subject = "SciTechDaily article export"
    props.comments = url
    target.parent.mkdir(parents=True, exist_ok=True)  # create the folder only now
    document.save(target)
    return image_count, title, category


def main() -> int:
    args = parse_args()
    if args.start > args.end:
        raise SystemExit("--start must not be later than --end")
    output = args.output or Path(f"SciTechDaily_{args.start.isoformat()}_to_{args.end.isoformat()}")
    output.mkdir(parents=True, exist_ok=True)
    session = make_session()
    entries = archive_articles(session, args.start, args.end)
    print(f"Found {len(entries)} articles from {args.start} through {args.end}")
    used: set[str] = set()
    failures: list[tuple[str, str]] = []
    for index, entry in enumerate(entries, 1):
        path = output / safe_filename(entry[1], used, output)
        try:
            count, _, _ = write_article(session, entry, path)
            print(f"[{index}/{len(entries)}] {path.name} ({count} images)")
        except Exception as exc:
            failures.append((entry[2], str(exc)))
            print(f"[{index}/{len(entries)}] FAILED {entry[2]}: {exc}")
        time.sleep(args.delay)
    if failures:
        (output / "failures.txt").write_text("\n".join(f"{url}\t{error}" for url, error in failures), encoding="utf-8")
    print(f"Completed: {len(entries) - len(failures)} documents, {len(failures)} failures in {output.resolve()}")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
