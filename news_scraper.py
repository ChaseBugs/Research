"""Multi-site science and technology news exporter."""

from __future__ import annotations

import argparse
import os
import re
import sys
import time
from dataclasses import dataclass, field as dataclass_field
from datetime import date, datetime
from email.utils import parsedate_to_datetime
from io import BytesIO
from pathlib import Path
from typing import Callable
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image

import scitechdaily_scraper
from activity_log import ActivityLog
from article_filter import KeywordFilter
from sites import FIELDS, SITES, Site, folder_label, sites_in_field


Log = Callable[[str], None]
# Returns True if the article is substantive-technical, False if news-only.
Classifier = Callable[[str, str], bool]
NEWS_ONLY_DIR = "보도기사"
INVALID_FILENAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
DEFAULT_LIMIT = 50
MAX_PATH_BUDGET = 250  # Windows MAX_PATH is 260; leave headroom


def scrape_output_folder(output_root: Path, site: Site,
                         collection_date: date | None = None) -> Path:
    """Daily DPRK Korean-language folder for one source: `MM-DD/<category>_<sub>(문서-영문)`."""
    day = collection_date or date.today()
    return output_root / day.strftime("%m-%d") / folder_label(site)
BODY_TAGS = ["figure", "p", "h2", "h3", "h4", "h5", "blockquote", "ul", "ol"]
JUNK_CLASS_WORDS = ("newsletter", "related", "advert", "social", "share", "author-bio")


def _classify(classify: Classifier | None, title: str, body: str) -> str:
    """'technical' or 'news'. No classifier => everything is technical."""
    if classify is None:
        return "technical"
    return "technical" if classify(title, body) else "news"


class ArticleSkipped(Exception):
    """Article deliberately not saved (excluded keyword or already collected)."""

    def __init__(self, reason: str) -> None:
        super().__init__(reason)
        self.reason = reason


@dataclass
class ScrapeResult:
    folder: Path
    saved: int = 0
    failed: int = 0
    skipped: int = 0
    news: int = 0        # saved but classified news-only (in the 보도기사 subfolder)

    @property
    def found(self) -> int:
        return self.saved + self.failed + self.skipped


@dataclass
class FieldResult:
    folder: Path
    per_site: dict[str, ScrapeResult] = dataclass_field(default_factory=dict)
    site_errors: dict[str, str] = dataclass_field(default_factory=dict)

    @property
    def saved(self) -> int:
        return sum(r.saved for r in self.per_site.values())

    @property
    def failed(self) -> int:
        return sum(r.failed for r in self.per_site.values())

    @property
    def skipped(self) -> int:
        return sum(r.skipped for r in self.per_site.values())

    @property
    def news(self) -> int:
        return sum(r.news for r in self.per_site.values())


def _text(tag: Tag | None) -> str:
    return " ".join(tag.get_text(" ", strip=True).split()) if tag else ""


def feed_articles(session: requests.Session, site: Site, start: date, end: date, log: Log) -> list[tuple[date, str, str]]:
    assert site.feed
    response = session.get(site.feed, timeout=60)
    response.raise_for_status()
    soup = BeautifulSoup(response.content, "xml")
    entries: dict[str, tuple[date, str, str]] = {}
    all_dates: list[date] = []
    for item in soup.find_all("item") or soup.find_all("entry"):
        title = _text(item.find("title"))
        link = _text(item.find("link")) or _text(item.find("guid"))
        if not link.startswith("http") and item.find("link") is not None:
            link = item.find("link").get("href", "") or link
        raw_date = (_text(item.find("pubDate")) or _text(item.find("published"))
                    or _text(item.find("updated")) or _text(item.find("date")))
        if not title or not link or not raw_date:
            continue
        try:
            published = parsedate_to_datetime(raw_date).date()
        except (TypeError, ValueError):
            try:
                published = datetime.fromisoformat(raw_date.replace("Z", "+00:00")).date()
            except ValueError:
                continue
        all_dates.append(published)
        if start <= published <= end:
            entries[link] = (published, title, link)
    if all_dates and start < min(all_dates):
        log(f"Warning: {site.name}'s oldest RSS article is {min(all_dates)}.")
    return sorted(entries.values(), key=lambda row: (row[0], row[1]), reverse=True)


def _image_url(img: Tag, page_url: str) -> str | None:
    value = img.get("data-lazy-src") or img.get("data-src") or img.get("src")
    if (not value or value.startswith("data:")) and img.get("srcset"):
        value = img["srcset"].split(",")[-1].strip().split()[0]
    return urljoin(page_url, value) if value and not value.startswith("data:") else None


def _add_image(document: Document, session: requests.Session, img: Tag, page_url: str, log: Log) -> bool:
    url = _image_url(img, page_url)
    if not url:
        return False
    try:
        response = session.get(url, timeout=60)
        response.raise_for_status()
        raw = BytesIO(response.content)
        with Image.open(raw) as parsed:
            parsed.load()
            converted = BytesIO()
            if "A" in parsed.getbands():
                parsed.save(converted, format="PNG")
            else:
                parsed.convert("RGB").save(converted, format="JPEG", quality=90)
        converted.seek(0)
        document.add_picture(converted, width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception as exc:
        log(f"  Image skipped: {url} ({exc})")
        return False


def _safe_filename(title: str, used: set[str], folder: Path,
                   existing_folders: tuple[Path, ...] = ()) -> str:
    """Build a unique .docx filename that fits inside the Windows path limit.

    Long headlines plus a nested field-run folder can push the full path past
    MAX_PATH (260), which surfaces as a confusing 'No such file or directory' on
    save, so the stem is budgeted against the folder it lands in.
    """
    stem = INVALID_FILENAME.sub("", title).rstrip(" .") or "Untitled article"
    budget = MAX_PATH_BUDGET - len(str(folder)) - len(".docx") - len(os.sep) - len(" (99)")
    budget = max(24, min(170, budget))
    stem = stem[:budget].rstrip(" .") or "Untitled article"
    candidate, number = stem, 2
    folders = (folder, *existing_folders)
    while (candidate.casefold() in used
           or any((directory / f"{candidate}.docx").exists() for directory in folders)):
        suffix = f" ({number})"
        candidate, number = stem[:budget - len(suffix)].rstrip(" .") + suffix, number + 1
    used.add(candidate.casefold())
    return candidate + ".docx"


def _screen_article(title: str, body: str, url: str,
                    keyword_filter: KeywordFilter | None,
                    dedup: ActivityLog | None) -> None:
    """Raise ArticleSkipped if the article is excluded or already collected.

    Runs on parsed text before the document is built, so a rejected article never
    costs an image download. `dedup`, when given, is the global ledger; a title
    seen in any earlier run (any day) counts as a duplicate.
    """
    if keyword_filter:
        hit = keyword_filter.match(title, body)
        if hit:
            raise ArticleSkipped(f"Excluded keyword '{hit}'")
    if dedup is not None and dedup.seen_title("scrape", title):
        raise ArticleSkipped("Title already collected")


def write_generic_article(session: requests.Session, site: Site, entry: tuple[date, str, str],
                          target: Path, log: Log,
                          keyword_filter: KeywordFilter | None = None,
                          dedup: ActivityLog | None = None,
                          classify: Classifier | None = None,
                          news_dir: Path | None = None) -> tuple[int, str, str]:
    _published, listed_title, url = entry
    response = session.get(url, timeout=60)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    content = next((soup.select_one(selector) for selector in site.content_selectors if soup.select_one(selector)), None)
    if not content:
        raise RuntimeError("Could not find the article body")
    title_tag = next((soup.select_one(selector) for selector in site.title_selectors if soup.select_one(selector)), None)
    title = _text(title_tag) or listed_title

    body_text = _text(content)
    _screen_article(title, body_text, url, keyword_filter, dedup)
    category = _classify(classify, title, body_text)
    if category == "news" and news_dir is not None:
        target = news_dir / target.name

    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading(title, level=0)

    images = 0
    emitted: set[str] = set()
    for selector in site.hero_selectors:
        hero = soup.select_one(selector)
        if not hero:
            continue
        for img in hero.find_all("img"):
            url_value = _image_url(img, url)
            if url_value and url_value not in emitted:
                images += int(_add_image(document, session, img, url, log))
                emitted.add(url_value)
        caption = hero.find("figcaption")
        if caption and _text(caption):
            paragraph = document.add_paragraph(_text(caption))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].italic = True
    for node in content.find_all(BODY_TAGS, recursive=True):
        if node.find_parent(BODY_TAGS) is not None and node.find_parent(BODY_TAGS) != content:
            continue
        classes = " ".join(node.get("class", [])).lower()
        if any(word in classes for word in JUNK_CLASS_WORDS):
            continue
        for img in node.find_all("img"):
            url_value = _image_url(img, url)
            if url_value and url_value not in emitted:
                images += int(_add_image(document, session, img, url, log))
                emitted.add(url_value)
        text = _text(node)
        if not text:
            continue
        if node.name == "figure":
            caption = node.find("figcaption")
            if caption and _text(caption):
                paragraph = document.add_paragraph(_text(caption))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].italic = True
        elif node.name in {"h2", "h3", "h4", "h5"}:
            document.add_heading(text, level=1 if node.name in {"h2", "h3"} else 2)
        elif node.name == "blockquote":
            document.add_paragraph(text, style="Quote")
        elif node.name in {"ul", "ol"}:
            style = "List Bullet" if node.name == "ul" else "List Number"
            for item in node.find_all("li", recursive=False):
                document.add_paragraph(_text(item), style=style)
        else:
            document.add_paragraph(text)
    document.core_properties.title = title
    document.core_properties.subject = f"{site.name} article export"
    document.core_properties.comments = url
    target.parent.mkdir(parents=True, exist_ok=True)  # create the folder only now
    document.save(target)
    return images, title, category


def scrape(site_name: str, start: date, end: date, output_root: Path, log: Log = print,
           delay: float = 0.2, limit: int | None = DEFAULT_LIMIT,
           keyword_filter: KeywordFilter | None = None,
           ledger: ActivityLog | None = None, skip_duplicates: bool = True,
           output: Path | None = None,
           classifier: Classifier | None = None,
           write_reports: bool = True) -> ScrapeResult:
    """Collect one site into the current day's Korean-language field folder.

    `output` overrides the destination folder, used by field runs so every
    source writes directly into the same field folder. Every saved article is recorded in the
    global `ledger`; when `skip_duplicates` is true, articles already in the ledger
    (from any earlier run, any day) are skipped by URL before fetching and by title
    after parsing. When `classifier` is given, articles it judges news-only are
    saved to a `보도기사/` subfolder instead of the main folder.
    """
    if site_name not in SITES:
        raise ValueError(f"Unsupported site: {site_name}")
    if start > end:
        raise ValueError("Start date must not be after end date")
    site = SITES[site_name]
    session = scitechdaily_scraper.make_session()
    if site.feed is None:
        entries = scitechdaily_scraper.archive_articles(session, start, end)
    else:
        entries = feed_articles(session, site, start, end, log)

    if output is None:
        output = scrape_output_folder(output_root, site)
    # Folders are created lazily, only when an article is actually saved into them
    # (see the writers), so a run with no new articles leaves no empty folder.
    news_dir = output / NEWS_ONLY_DIR if classifier is not None else None

    dedup = ledger if (ledger is not None and skip_duplicates) else None
    total_found = len(entries)
    if dedup is not None:
        entries = [row for row in entries if not dedup.seen_key("scrape", row[2])]
    already = total_found - len(entries)
    if limit:
        entries = entries[:limit]

    note = f"{site.name}: found {total_found} articles"
    if already:
        note += f", excluded {already} from collection history"
    if limit and total_found - already > limit:
        note += f", processing only the latest {limit}"
    log(note)

    result = ScrapeResult(folder=output, skipped=already)
    used: set[str] = set()
    failures: list[str] = []
    skips: list[str] = []
    for index, entry in enumerate(entries, 1):
        published, listed_title, url = entry
        target = output / _safe_filename(
            listed_title,
            used,
            output,
            (news_dir,) if news_dir is not None else (),
        )
        try:
            if site.feed is None:
                images, title, category = scitechdaily_scraper.write_article(
                    session, entry, target,
                    screen=lambda t, b: _screen_article(t, b, url, keyword_filter, dedup),
                    classify=classifier, news_dir=news_dir,
                )
            else:
                images, title, category = write_generic_article(
                    session, site, entry, target, log, keyword_filter, dedup,
                    classify=classifier, news_dir=news_dir)
            saved_dir = news_dir if (category == "news" and news_dir is not None) else output
            tag = " [News]" if category == "news" else ""
            log(f"[{index}/{len(entries)}]{tag} {target.name} ({images} images)")
            result.saved += 1
            if category == "news":
                result.news += 1
            if ledger is not None:
                ledger.record("scrape", title, key=url, path=str(saved_dir / target.name),
                              source=site.name, keyword=site.field)
        except ArticleSkipped as skip:
            skips.append(f"{url}\t{skip.reason}")
            log(f"[{index}/{len(entries)}] Skipped ({skip.reason}): {listed_title}")
            result.skipped += 1
            target.unlink(missing_ok=True)
        except Exception as exc:
            failures.append(f"{url}\t{exc}")
            log(f"[{index}/{len(entries)}] Failed: {listed_title} - {exc}")
            result.failed += 1

    # Only leave report files when the folder exists (i.e. at least one .docx was
    # saved); a run that produced no documents leaves nothing behind.
    if write_reports and output.exists():
        _write_report(output / "failures.txt", failures)
        _write_report(output / "_skipped.txt", skips)
    tech = result.saved - result.news
    news_note = f" (technology {tech}, news {result.news})" if classifier is not None else ""
    log(f"Complete: saved {result.saved}{news_note}, skipped {result.skipped}, failed {result.failed}")
    return result


def scrape_field(field_key: str, start: date, end: date, output_root: Path, log: Log = print,
                 delay: float = 0.2, limit: int | None = DEFAULT_LIMIT,
                 keyword_filter: KeywordFilter | None = None,
                 ledger: ActivityLog | None = None, skip_duplicates: bool = True,
                 classifier: Classifier | None = None) -> FieldResult:
    """Collect every site in a field into one daily Korean-language folder.

    A site that fails outright is recorded and the run continues to the next.
    """
    if field_key not in FIELDS:
        raise ValueError(f"Unsupported field: {field_key}")
    members = sites_in_field(field_key)
    if not members:
        raise ValueError(f"No sites registered in field: {field_key}")

    # Each source writes to its own <category>_<sub-category>(문서-영문) folder under
    # the day folder; the sub-category differs per source, so there is no single
    # shared field folder any more.
    day_folder = output_root / date.today().strftime("%m-%d")
    day_folder.mkdir(parents=True, exist_ok=True)
    result = FieldResult(folder=day_folder)
    log(f"=== {FIELDS[field_key]} / {len(members)} sites ===")

    for position, site in enumerate(members, 1):
        log(f"\n--- [{position}/{len(members)}] {site.name} -> {folder_label(site)} ---")
        try:
            result.per_site[site.name] = scrape(
                site.name, start, end, output_root, log, delay, limit,
                keyword_filter, ledger, skip_duplicates,
                classifier=classifier,
                write_reports=False,
            )
        except Exception as exc:
            result.site_errors[site.name] = str(exc)
            log(f"!!! {site.name}: all requests failed: {exc}")
        time.sleep(delay)

    log(f"\n=== Field complete: saved {result.saved}, skipped {result.skipped}, "
        f"failed {result.failed}, site errors {len(result.site_errors)} ===")
    return result


def _write_report(path: Path, lines: list[str]) -> None:
    if lines:
        path.write_text("\n".join(lines), encoding="utf-8")
    elif path.exists():
        path.unlink()


def _print_catalogue() -> None:
    for key, label in FIELDS.items():
        members = sites_in_field(key)
        print(f"\n{key}  ({label})  — {len(members)} sites")
        for site in members:
            mark = "  [low volume]" if site.low_volume else ""
            print(f"    {site.name}{mark}")
    print(f"\nTotal: {len(SITES)} sites / {len(FIELDS)} fields")


def _use_utf8_console() -> None:
    """Korean log output crashes on a cp1252 console without this."""
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")


def main() -> int:
    _use_utf8_console()
    parser = argparse.ArgumentParser(description=__doc__)
    target = parser.add_mutually_exclusive_group()
    target.add_argument("--site", choices=SITES)
    target.add_argument("--field", choices=FIELDS)
    target.add_argument("--list", action="store_true", help="Print registered fields and sites")
    parser.add_argument("--start", type=date.fromisoformat)
    parser.add_argument("--end", type=date.fromisoformat)
    parser.add_argument("--output", type=Path, default=Path.cwd())
    parser.add_argument("--limit", type=int, default=DEFAULT_LIMIT,
                        help="Maximum articles per site (0 = unlimited)")
    parser.add_argument("--keywords", type=Path, default=None,
                        help="Excluded-keyword file (default: exclude_keywords.txt)")
    parser.add_argument("--no-history", action="store_true",
                        help="Do not skip duplicates")
    args = parser.parse_args()

    if args.list:
        _print_catalogue()
        return 0
    if not args.site and not args.field:
        parser.error("Specify either --site or --field (use --list to see options)")
    if not args.start or not args.end:
        parser.error("Both --start and --end are required")

    keyword_filter = KeywordFilter.from_file(args.keywords)
    if keyword_filter:
        print(f"Applied {len(keyword_filter)} excluded keywords")
    ledger = ActivityLog.load(args.output)
    skip_duplicates = not args.no_history
    if skip_duplicates:
        print(f"Collection history: {len(ledger.entries())} entries (SQLite)")
    limit = args.limit or None

    if args.field:
        result = scrape_field(args.field, args.start, args.end, args.output,
                              limit=limit, keyword_filter=keyword_filter,
                              ledger=ledger, skip_duplicates=skip_duplicates)
        failures = result.failed + len(result.site_errors)
    else:
        result = scrape(args.site, args.start, args.end, args.output,
                        limit=limit, keyword_filter=keyword_filter,
                        ledger=ledger, skip_duplicates=skip_duplicates)
        failures = result.failed
    return int(bool(failures))


if __name__ == "__main__":
    raise SystemExit(main())
