"""Turn a keyword into a list of technology topics, then into illustrated DOCX guides.

Two stages, matching the Generate tab's two buttons:

1. ``generate_topics(keyword, n)`` asks the model for *n* concrete technology
   titles related to the keyword (e.g. "mining town modernization" -> 100 mining-modernization
   technologies), returned as a clean de-duplicated list.
2. ``generate_document(topic, keyword, ...)`` writes one beginner-oriented guide
   per topic -- what it is, how it works, how and where to use it, how to
   implement it -- as a DOCX with AI-generated illustrations.

All model output is English by instruction; Korean is intentionally avoided so the
documents match the workspace's English corpus.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image

from ai_client import AIClient, AIError

INVALID_FILENAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_LIST_PREFIX = re.compile(r"^\s*(?:\d+[.)]|[-*•])\s+")
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_HEADING = re.compile(r"^(#{1,4})\s+(.*)$")

TOPIC_SYSTEM = "You are a technical research assistant. Answer in English only."
DOC_SYSTEM = (
    "You are a technical writer producing clear, practical documentation for "
    "beginners. Answer in English only. Be concrete and specific."
)

DOC_SECTIONS = (
    "Overview (what this technology is)",
    "How It Works (the underlying principle)",
    "How to Use It (practical operation)",
    "Where It Is Used (real applications and settings)",
    "Implementation Steps (so a beginner can start building or applying it)",
    "Key Considerations (limitations, safety, cost)",
)


def _clean_topic(line: str) -> str:
    text = _LIST_PREFIX.sub("", line).strip()
    text = _BOLD.sub(r"\1", text).strip(" .\t")
    return text


def generate_topics(ai: AIClient, keyword: str, count: int = 100,
                    log=lambda _m: None) -> list[str]:
    """Return up to `count` distinct technology titles related to `keyword`.

    The model rarely returns exactly `count` usable lines, so this requests,
    parses and de-duplicates in rounds until it has enough or stops making
    progress.
    """
    topics: list[str] = []
    seen: set[str] = set()
    rounds = 0
    while len(topics) < count and rounds < 6:
        rounds += 1
        remaining = count - len(topics)
        avoid = ""
        if topics:
            avoid = (" Do NOT repeat any of these already-listed items:\n"
                     + "; ".join(topics[-40:]))
        prompt = (
            f"The goal is technical documentation about: \"{keyword}\".\n"
            f"List {remaining} distinct, specific technologies, engineering "
            f"methods, systems or techniques related to \"{keyword}\". "
            f"Each must be a concrete technology name usable as a document title "
            f"(a noun phrase, not a sentence). "
            f"Output ONLY a numbered list, one item per line, in English.{avoid}"
        )
        try:
            reply = ai.chat(prompt, system=TOPIC_SYSTEM)
        except AIError as exc:
            log(f"Topic generation failed: {exc}")
            break
        added = 0
        for line in reply.splitlines():
            topic = _clean_topic(line)
            if len(topic) < 3 or len(topic) > 160:
                continue
            key = topic.casefold()
            if key in seen:
                continue
            seen.add(key)
            topics.append(topic)
            added += 1
            if len(topics) >= count:
                break
        log(f"[Round {rounds}] Added {added} (total {len(topics)}/{count})")
        if added == 0:
            break
    return topics[:count]


# --------------------------------------------------------------------- document

def _add_markdown(document: Document, text: str) -> None:
    """Render the model's lightly-marked-up prose into DOCX blocks."""
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        heading = _HEADING.match(line)
        if heading:
            level = min(len(heading.group(1)), 3)
            document.add_heading(_BOLD.sub(r"\1", heading.group(2)).strip(), level=level)
            continue
        if _LIST_PREFIX.match(line):
            body = _BOLD.sub(r"\1", _LIST_PREFIX.sub("", line)).strip()
            numbered = bool(re.match(r"^\s*\d+[.)]", line))
            document.add_paragraph(body, style="List Number" if numbered else "List Bullet")
            continue
        document.add_paragraph(_BOLD.sub(r"\1", line).strip())


def _add_image(document: Document, ai: AIClient, prompt: str, caption: str) -> bool:
    buffer = ai.try_image(prompt, width=768, height=512)
    if buffer is None:
        return False
    try:
        converted = _to_jpeg(buffer)
        document.add_picture(converted, width=Inches(6.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = document.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        return True
    except Exception:
        return False


def _to_jpeg(buffer):
    from io import BytesIO
    with Image.open(buffer) as parsed:
        parsed.load()
        out = BytesIO()
        parsed.convert("RGB").save(out, format="JPEG", quality=88)
    out.seek(0)
    return out


def safe_filename(title: str, used: set[str], folder: Path) -> str:
    stem = INVALID_FILENAME.sub("", title).rstrip(" .") or "Untitled"
    budget = 250 - len(str(folder)) - len(".docx") - 6
    budget = max(24, min(170, budget))
    stem = stem[:budget].rstrip(" .") or "Untitled"
    candidate, n = stem, 2
    while candidate.casefold() in used:
        suffix = f" ({n})"
        candidate, n = stem[:budget - len(suffix)].rstrip(" .") + suffix, n + 1
    used.add(candidate.casefold())
    return candidate + ".docx"


def generate_document(ai: AIClient, topic: str, keyword: str, target: Path,
                      with_images: bool = True, log=lambda _m: None) -> tuple[int, str]:
    """Write one illustrated technical guide to `target`. Returns (image count, title)."""
    prompt = (
        f"Write a detailed, beginner-friendly technical guide about "
        f"\"{topic}\" in the context of \"{keyword}\".\n"
        f"Use these sections, each introduced by a Markdown '## ' heading:\n"
        + "\n".join(f"- {s}" for s in DOC_SECTIONS)
        + "\nUse '- ' for bullet points and '1. ' for ordered steps. "
        f"Write so a beginner can understand and start applying the technology. "
        f"Answer in English only."
    )
    body = ai.chat(prompt, system=DOC_SYSTEM)

    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading(topic, level=0)

    images = 0
    if with_images:
        if _add_image(document, ai,
                      f"technical schematic diagram illustrating {topic}, {keyword}, "
                      f"labeled, engineering blueprint style",
                      f"Figure 1. Conceptual illustration of {topic}."):
            images += 1

    _add_markdown(document, body)

    if with_images:
        if _add_image(document, ai,
                      f"realistic photo of {topic} in use, {keyword}, industrial setting",
                      f"Figure 2. {topic} in a practical setting."):
            images += 1

    document.core_properties.title = topic
    document.core_properties.subject = f"Generated technical guide: {keyword}"
    document.core_properties.comments = f"AI-generated document for '{keyword}'"
    document.save(target)
    log(f"Saved: {target.name} ({images} images)")
    return images, topic
